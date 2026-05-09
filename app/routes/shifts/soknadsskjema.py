import logging
import os
import tempfile
from datetime import date

from flask import flash, render_template, request, send_file
from flask_login import current_user, login_required

from app.database import get_db_session
from app.models import DBUser, SoknadsskjemaChoice
from app.routes.shifts import shifts
from app.utils import db_utils
from app.utils.turnus_helpers import get_user_turnus_set

logger = logging.getLogger(__name__)


def _set_table_col_widths(table, col_widths_dxa, add_borders=True):
    """Set explicit tblGrid and per-cell tcW widths for cross-app compatibility.

    IMPORTANT: call this AFTER all cell merges so that gridSpan is already set
    and tcW can reflect the correct merged width. Using row._tr.iterchildren()
    instead of row.cells avoids the python-docx behaviour of returning the same
    merged cell object multiple times (once per logical column it spans), which
    would produce wrong tcW values and cause the sum of per-row tcW to diverge
    from tblW — a mismatch OpenOffice treats as a fatal table error.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))

    # Total table width
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(col_widths_dxa)))
    tblW.set(qn("w:type"), "dxa")

    # ── tblBorders (schema pos 11): explicit borders for cross-app compat ────────────
    # OO/LO does not always apply style-level borders from tblStyle; inline tblBorders
    # are always honoured regardless of style resolution.
    # When add_borders=False the caller manages borders at the cell level instead.
    if add_borders and tblPr.find(qn("w:tblBorders")) is None:
        tblBorders = OxmlElement("w:tblBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "4")
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "auto")
            tblBorders.append(b)
        # Insert before tblLook (pos 15) to respect OOXML schema order
        tblLook = tblPr.find(qn("w:tblLook"))
        if tblLook is not None:
            tblLook.addprevious(tblBorders)
        else:
            tblPr.append(tblBorders)

    # ── tblLayout fixed (schema pos 13): after tblBorders, BEFORE tblLook ────────────
    # OO/LO requires this declaration to treat tblW/tcW as hard constraints.
    # CRITICAL: must be inserted before tblLook (pos 15); appending after tblLook
    # violates schema order and causes OO to silently discard the element.
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblLook = tblPr.find(qn("w:tblLook"))
        if tblLook is not None:
            tblLook.addprevious(tblLayout)
        else:
            tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")

    # Replace tblGrid so renderers know exact column widths
    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)
    new_grid = OxmlElement("w:tblGrid")
    for w in col_widths_dxa:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        new_grid.append(gc)
    tblPr.addnext(new_grid)

    # Per-cell explicit widths — iterate actual <w:tc> elements, not row.cells.
    # row.cells expands merged cells into one entry per logical column, so a cell
    # with gridSpan=2 appears twice; iterchildren gives exactly one entry per
    # physical cell, matching the gridSpan already written by merge().
    for row in table.rows:
        ci = 0
        for tc in row._tr.iterchildren(qn("w:tc")):
            if ci >= len(col_widths_dxa):
                break
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                tc.insert(0, tcPr)
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.insert(0, tcW)
            grid_span = tcPr.find(qn("w:gridSpan"))
            span = int(grid_span.get(qn("w:val"), 1)) if grid_span is not None else 1
            tcW.set(qn("w:w"), str(sum(col_widths_dxa[ci : ci + span])))
            tcW.set(qn("w:type"), "dxa")
            ci += span


def _add_cell_border(cell):
    """Add a single-line border on all four sides of one table cell (tcBorders)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "595959")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _arial(run, size_pt, bold=False):
    from docx.shared import Pt

    run.font.name = "Arial"
    run.font.size = Pt(size_pt)
    if bold:
        run.font.bold = True


def _get_soknadsskjema_choices(user_id, turnus_set_id):
    """Return {shift_title: {linje_135, linje_246, h_dag, linjeprioritering}} from DB."""
    db_session = get_db_session()
    try:
        rows = (
            db_session.query(SoknadsskjemaChoice)
            .filter_by(user_id=user_id, turnus_set_id=turnus_set_id)
            .all()
        )
        return {
            r.shift_title: {
                "linje_135": bool(r.linje_135),
                "linje_246": bool(r.linje_246),
                "h_dag": bool(r.h_dag),
                "linjeprioritering": r.linjeprioritering or "",
            }
            for r in rows
        }
    except Exception as e:
        logger.error("_get_soknadsskjema_choices error: %s", e)
        return {}
    finally:
        db_session.close()


def _build_soknadsskjema_doc(
    dato, rullenr_og_navn, stasjoneringssted, kommentarer, favorites, choices=None
):
    """Generate søknadsskjema from scratch matching the original form layout.

    Layout (top to bottom):
      1. Header table: left cell = title + "Unngå stifter..." instructions;
                       right cell = bordered "Fylles ut av Forening" box
      2. Personal info table (Dato / Rullenr. / Stasjoneringssted / Kommentarer)
      3. Instruction text about Linje (with bold + partial underlines)
      4. Alt table (71 rows, merges done BEFORE _set_table_col_widths)

    The merge-before-widths order is critical: python-docx's row.cells expands
    merged cells (one entry per logical column), so calling _set_table_col_widths
    before merging leaves merged cells with single-column tcW. The resulting
    tcW-sum < tblW mismatch is treated as a fatal table error by OpenOffice.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = Document()

    # Page: A4, margins matching original
    section = doc.sections[0]
    section.page_width = Pt(595)
    section.page_height = Pt(842)
    section.left_margin = Pt(71)
    section.right_margin = Pt(71)
    section.top_margin = Pt(27)
    section.bottom_margin = Pt(35)
    # header_distance must be < top_margin (27 pt); default is 36 pt which would
    # place the header below the body text start. 12 pt → 15 pt clearance to body.
    section.header_distance = Pt(12)

    # Pages 2+: rullenr og navn right-aligned. Page 1 gets no header (titlePg).
    section.different_first_page_header_footer = True
    hdr_para = section.header.paragraphs[0]
    hdr_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _arial(hdr_para.add_run(rullenr_og_navn), 9)

    # Strip python-docx default 8 pt space-after + 1.15× line spacing from Normal
    normal = doc.styles["Normal"]
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    def _p(text="", size_pt=11, bold=False):
        para = doc.add_paragraph()
        if text:
            _arial(para.add_run(text), size_pt, bold)
        return para

    def _mixed(size_pt=11, *parts):
        """Paragraph with mixed runs: each part is (text, bold, underline)."""
        para = doc.add_paragraph()
        for text, bold, underline in parts:
            r = para.add_run(text)
            _arial(r, size_pt, bold)
            if underline:
                r.font.underline = True
        return para

    # ── Header table: title + top instructions (left cell) + forening box (right cell) ──
    # The box spans the full height of both the title and the instruction lines,
    # matching the screenshot layout where the box sits alongside all of that content.
    HEADER_COL_WIDTHS = [7020, 2040]  # sum = 9060 dxa; right col ≈ 141px
    header_tbl = doc.add_table(rows=1, cols=2)
    _set_table_col_widths(header_tbl, HEADER_COL_WIDTHS, add_borders=False)

    # Left cell: title (16 pt bold) then a gap then the two instruction lines (11 pt)
    left = header_tbl.cell(0, 0)
    _arial(left.paragraphs[0].add_run("Søknad turplassering for"), 16, bold=True)
    _arial(left.add_paragraph().add_run("Lokomotivpersonalet"), 16, bold=True)
    left.add_paragraph()  # visual gap between title and instructions
    _arial(left.add_paragraph().add_run("Unngå stifter og tape"), 11)
    _arial(
        left.add_paragraph().add_run(
            "Bruk helst ensidig og merk hver ark med navn og rullenr"
        ),
        11,
    )

    # Right cell: bordered box — "Fylles ut av Forening" at top, "Ansiennitet:" near bottom
    right = header_tbl.cell(0, 1)
    _add_cell_border(right)
    gray = RGBColor(0x55, 0x55, 0x55)
    r1 = right.paragraphs[0].add_run("Fylles ut av Forening")
    _arial(r1, 8)
    r1.font.color.rgb = gray
    for _ in range(3):  # spacers so "Ansiennitet:" sits near the bottom
        right.add_paragraph()
    r3 = right.add_paragraph().add_run("Ansiennitet:")
    _arial(r3, 8)
    r3.font.color.rgb = gray

    # ── Personal info table ──
    _p()
    P_COL_WIDTHS = [2856, 6204]  # sum = 9060 dxa
    p_tbl = doc.add_table(rows=4, cols=2, style="Table Grid")
    _set_table_col_widths(p_tbl, P_COL_WIDTHS)
    for i, (label, value) in enumerate(
        [
            ("Dato", dato),
            ("Rullenr. og navn", rullenr_og_navn),
            ("Stasjoneringsted", stasjoneringssted),
            ("Evt. kommentarer", kommentarer),
        ]
    ):
        _arial(p_tbl.rows[i].cells[0].paragraphs[0].add_run(label), 11, bold=True)
        _arial(p_tbl.rows[i].cells[1].paragraphs[0].add_run(value), 11)

    # ── Middle instruction (comes AFTER personal info table, matching original layout) ──
    _p()
    _mixed(
        11,
        ("Linje er ", True, False),
        ("uten", True, True),
        (" betydning:", True, False),
    )
    _p("Fyll kun ut kolonne 1.")
    _p("Du plasseres i vilkårlig valgt linje.")
    _mixed(
        11,
        ("Kun ", True, False),
        ("helg", True, True),
        (" er av betydning:", True, False),
    )
    _p("Fyll ut kolonne 1 og 2")
    _p(
        "Du plasseres i vilkårlig valgt linje innenfor din helg (Linje 1,3,5 eller 2,4,6)"
    )
    _mixed(11, ("Linje", True, True), (" er av betydning", True, False))
    _p("Fyll ut kolonne 1 og 3")
    _mixed(
        11,
        ("Skriv linjer i ", False, False),
        ("prioritert rekkefølge", False, True),
        (" i kolonne 3. Du søker kun de linjene som er ført opp.", False, False),
    )

    # ── Alt table ──
    # Widths scaled from original proportions to text area (9060 dxa = 595pt − 2×71pt).
    COL_WIDTHS = [805, 2500, 1000, 1000, 2307, 1448]  # sum = 9060 dxa

    alt_tbl = doc.add_table(rows=3 + 71, cols=6, style="Table Grid")

    def _cell(r, c, text, size_pt=10, bold=False):
        _arial(alt_tbl.cell(r, c).paragraphs[0].add_run(text), size_pt, bold)

    # Merges FIRST so gridSpan is set before _set_table_col_widths reads it
    alt_tbl.cell(0, 0).merge(alt_tbl.cell(0, 1))
    alt_tbl.cell(0, 2).merge(alt_tbl.cell(0, 3))
    alt_tbl.cell(1, 0).merge(alt_tbl.cell(1, 1))
    alt_tbl.cell(1, 2).merge(alt_tbl.cell(1, 3))

    _set_table_col_widths(alt_tbl, COL_WIDTHS)

    # Header row 0: Kolonne labels
    _cell(0, 0, "Kolonne 1", bold=True)
    _cell(0, 2, "Kolonne 2", bold=True)
    _cell(0, 4, "Kolonne 3", bold=True)
    _cell(0, 5, "Kolonne 4", bold=True)

    # Header row 1: column descriptions (matching original text)
    _cell(1, 0, "Tur\nnummer:", size_pt=8)
    _cell(
        1,
        2,
        "Ønsker en av følgende linjer:\n(Sett X)\n(Ingen prioritering blant disse)",
        size_pt=8,
    )
    _cell(
        1,
        4,
        "Linjeprioritering\n(Skriv inn de linjene du ønsker i prioritert rekkefølge)",
        size_pt=8,
    )
    _cell(1, 5, "H-dag\n(Skriv J for jobb.\nBlankt felt gir fri)", size_pt=8)

    # Header row 2: sub-column labels + down arrow indicating Tur nummer entry column
    _cell(2, 1, "↓", size_pt=8)
    _cell(2, 2, "Linje 1,3,5", size_pt=8)
    _cell(2, 3, "Linje 2,4,6", size_pt=8)

    # Data rows: Alt.1 – Alt.71
    for i in range(71):
        _cell(3 + i, 0, f"Alt.{i + 1}")
        if i < len(favorites):
            name = favorites[i]
            _cell(3 + i, 1, name.replace("_", " "))
            if choices:
                c = choices.get(name, {})
                if c.get("linje_135"):
                    _cell(3 + i, 2, "X")
                if c.get("linje_246"):
                    _cell(3 + i, 3, "X")
                if c.get("linjeprioritering"):
                    _cell(3 + i, 4, c["linjeprioritering"])
                if c.get("h_dag"):
                    _cell(3 + i, 5, "J")

    # OOXML requires the document body to end with a paragraph, not a table.
    # Some processors (including older OO/LO) misparse documents that end with a bare tbl.
    _p()

    return doc


def _build_soknadsskjema_pdf(
    dato, rullenr_og_navn, stasjoneringssted, kommentarer, favorites, choices=None
):
    """Generate søknadsskjema as PDF bytes matching the docx layout."""
    from io import BytesIO

    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    buf = BytesIO()
    PAGE_W, PAGE_H = 595, 842
    LM = RM = 71
    TW = PAGE_W - LM - RM  # 453 pt — same text width as docx
    doc = SimpleDocTemplate(
        buf,
        pagesize=(PAGE_W, PAGE_H),
        leftMargin=LM,
        rightMargin=RM,
        topMargin=27,
        bottomMargin=35,
    )

    def _ps(name, fn="Helvetica", fs=11, ld=14, **kw):
        return ParagraphStyle(
            name,
            fontName=fn,
            fontSize=fs,
            leading=ld,
            spaceAfter=0,
            spaceBefore=0,
            **kw,
        )

    body = _ps("ss_body", fs=10.5, ld=13)
    title = _ps("ss_title", fn="Helvetica-Bold", fs=16, ld=20)
    cell = _ps("ss_cell", fs=9, ld=11)
    cell_b = _ps("ss_cell_b", fn="Helvetica-Bold", fs=9, ld=11)
    alt_n = _ps("ss_alt_n", fs=9, ld=11, textColor=colors.HexColor("#666666"))
    box_s = _ps("ss_box", fs=8.5, ld=10.5, textColor=colors.HexColor("#555555"))

    BORDER = colors.HexColor("#444444")
    story = []

    # ── 1. Title + "Fylles ut av Forening" box ──
    hdr = Table(
        [
            [
                Paragraph("Søknad turplassering for<br/>Lokomotivpersonalet", title),
                Paragraph("Fylles ut av Forening<br/><br/>Ansiennitet:", box_s),
            ]
        ],
        colWidths=[TW - 130, 130],
    )
    hdr.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("BOX", (1, 0), (1, 0), 0.5, colors.HexColor("#555555")),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("LEFTPADDING", (0, 0), (0, 0), 0),
                ("RIGHTPADDING", (0, 0), (0, 0), 0),
                ("LEFTPADDING", (1, 0), (1, 0), 6),
                ("RIGHTPADDING", (1, 0), (1, 0), 6),
                ("TOPPADDING", (1, 0), (1, 0), 4),
                ("BOTTOMPADDING", (1, 0), (1, 0), 4),
            ]
        )
    )
    story += [hdr, Spacer(1, 6)]

    # ── 2. Top instruction ──
    story += [
        Paragraph("Unngå stifter og tape", body),
        Paragraph("Bruk helst ensidig og merk hver ark med navn og rullenr", body),
        Spacer(1, 8),
    ]

    # ── 3. Personal info table ──
    # Col widths: [2856, 6204] dxa / 20 = [142.8, 310.2] pt
    p_tbl = Table(
        [
            [Paragraph("<b>Dato</b>", body), Paragraph(dato or "", body)],
            [
                Paragraph("<b>Rullenr. og navn</b>", body),
                Paragraph(rullenr_og_navn or "", body),
            ],
            [
                Paragraph("<b>Stasjoneringsted</b>", body),
                Paragraph(stasjoneringssted or "", body),
            ],
            [
                Paragraph("<b>Evt. kommentarer</b>", body),
                Paragraph(kommentarer or "", body),
            ],
        ],
        colWidths=[143, 310],
    )
    p_tbl.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.5, BORDER),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story += [p_tbl, Spacer(1, 8)]

    # ── 4. Middle instruction ──
    story += [
        Paragraph("<b>Linje er <u>uten</u> betydning:</b>", body),
        Paragraph("Fyll kun ut kolonne 1.", body),
        Paragraph("Du plasseres i vilkårlig valgt linje.", body),
        Paragraph("<b>Kun <u>helg</u> er av betydning:</b>", body),
        Paragraph("Fyll ut kolonne 1 og 2", body),
        Paragraph(
            "Du plasseres i vilkårlig valgt linje innenfor din helg (Linje 1,3,5 eller 2,4,6)",
            body,
        ),
        Paragraph("<b><u>Linje</u> er av betydning</b>", body),
        Paragraph("Fyll ut kolonne 1 og 3", body),
        Paragraph(
            "Skriv linjer i <u>prioritert rekkefølge</u> i kolonne 3. Du søker kun de linjene som er ført opp.",
            body,
        ),
        Spacer(1, 6),
    ]

    # ── 5. Alt table ──
    # Col widths: [805, 2500, 1000, 1000, 2307, 1448] dxa / 20 → [40.25, 125, 50, 50, 115.35, 72.4] pt
    AW = [40.25, 125, 50, 50, 115.35, 72.4]
    cell_xs = _ps("ss_cell_xs", fs=7.5, ld=9.5)
    alt_rows = [
        # Row 0 — Kolonne headers; empty strings mark the spanned cells
        [
            Paragraph("Kolonne 1", cell_b),
            "",
            Paragraph("Kolonne 2", cell_b),
            "",
            Paragraph("Kolonne 3", cell_b),
            Paragraph("Kolonne 4", cell_b),
        ],
        # Row 1 — column descriptions
        [
            Paragraph("Tur<br/>nummer:", cell_xs),
            "",
            Paragraph(
                "Ønsker en av følgende linjer:<br/>(Sett X)<br/>(Ingen prioritering blant disse)",
                cell_xs,
            ),
            "",
            Paragraph(
                "Linjeprioritering<br/>(Skriv inn de linjene du ønsker i prioritert rekkefølge)",
                cell_xs,
            ),
            Paragraph(
                "H-dag<br/>(Skriv J for jobb.<br/>Blankt felt gir fri)",
                cell_xs,
            ),
        ],
        # Row 2 — sub-column labels
        [
            "",
            Paragraph("↓", cell_xs),
            Paragraph("Linje 1,3,5", cell_xs),
            Paragraph("Linje 2,4,6", cell_xs),
            "",
            "",
        ],
    ]
    for i in range(71):
        fav = favorites[i] if i < len(favorites) else ""
        c = choices.get(fav, {}) if (choices and fav) else {}
        alt_rows.append(
            [
                Paragraph(f"Alt.{i + 1}", alt_n),
                fav.replace("_", " "),
                "X" if c.get("linje_135") else "",
                "X" if c.get("linje_246") else "",
                c.get("linjeprioritering", ""),
                "J" if c.get("h_dag") else "",
            ]
        )

    alt_tbl = Table(alt_rows, colWidths=AW)
    alt_tbl.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.5, BORDER),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("LEADING", (0, 0), (-1, -1), 11),
                ("TOPPADDING", (0, 0), (-1, -1), 1),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
                ("LEFTPADDING", (0, 0), (-1, -1), 2),
                ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                # Header background + bold
                ("BACKGROUND", (0, 0), (-1, 2), colors.HexColor("#f2f2f2")),
                ("FONTNAME", (0, 0), (-1, 2), "Helvetica-Bold"),
                # Column merges: Kolonne 1 spans cols 0-1, Kolonne 2 spans cols 2-3 (rows 0 and 1)
                ("SPAN", (0, 0), (1, 0)),
                ("SPAN", (2, 0), (3, 0)),
                ("SPAN", (0, 1), (1, 1)),
                ("SPAN", (2, 1), (3, 1)),
                # Alignment
                ("ALIGN", (0, 0), (-1, 2), "CENTER"),
                ("ALIGN", (2, 3), (3, -1), "CENTER"),
                ("ALIGN", (5, 3), (5, -1), "CENTER"),
                # Tur name: bold
                ("FONTNAME", (1, 3), (1, -1), "Helvetica-Bold"),
            ]
        )
    )
    story.append(alt_tbl)

    def _later_pages(canvas, doc):
        """Draw rullenr og navn right-aligned at the top of pages 2+."""
        canvas.saveState()
        canvas.setFont("Helvetica", 9)
        canvas.drawRightString(PAGE_W - RM, PAGE_H - 25, rullenr_og_navn or "")
        canvas.restoreState()

    doc.build(story, onLaterPages=_later_pages)
    buf.seek(0)
    return buf


@shifts.route("/soknadsskjema", methods=["GET", "POST"])
@login_required
def soknadsskjema():
    user_turnus_set = get_user_turnus_set()
    turnus_set_id = user_turnus_set["id"] if user_turnus_set else None
    user_id = current_user.get_id()

    fav_order_lst = db_utils.get_favorite_lst(user_id, turnus_set_id)

    # Pre-populate personal info from DBUser
    db_session = get_db_session()
    try:
        db_user = db_session.query(DBUser).filter_by(id=user_id).first()
        user_name = (db_user.name or "") if db_user else ""
        user_rullenummer = (db_user.rullenummer or "") if db_user else ""
        user_stasjoneringssted = (db_user.stasjoneringssted or "") if db_user else ""
    finally:
        db_session.close()

    choices = (
        _get_soknadsskjema_choices(user_id, turnus_set_id) if turnus_set_id else {}
    )

    if request.method == "POST":
        dato = request.form.get("dato", "")
        rullenr_og_navn = request.form.get("rullenr_og_navn", "")
        stasjoneringssted = request.form.get("stasjoneringssted", "")
        kommentarer = request.form.get("kommentarer", "")
        fmt = request.form.get("format", "docx")

        year_id = user_turnus_set["year_identifier"] if user_turnus_set else "turnus"

        try:
            if fmt == "pdf":
                pdf_buf = _build_soknadsskjema_pdf(
                    dato,
                    rullenr_og_navn,
                    stasjoneringssted,
                    kommentarer,
                    fav_order_lst,
                    choices=choices,
                )
                return send_file(
                    pdf_buf,
                    as_attachment=True,
                    download_name=f"soknadsskjema_{year_id}.pdf",
                    mimetype="application/pdf",
                )
            else:
                doc = _build_soknadsskjema_doc(
                    dato,
                    rullenr_og_navn,
                    stasjoneringssted,
                    kommentarer,
                    fav_order_lst,
                    choices=choices,
                )
                filename = f"soknadsskjema_{year_id}.docx"
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                temp_file_path = temp_file.name
                temp_file.close()
                doc.save(temp_file_path)

                response = send_file(
                    temp_file_path,
                    as_attachment=True,
                    download_name=filename,
                    mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

                @response.call_on_close
                def cleanup():
                    if os.path.exists(temp_file_path):
                        os.unlink(temp_file_path)

                return response

        except Exception as e:
            logger.error("Error generating soknadsskjema (%s): %s", fmt, e)
            flash("Feil ved generering av søknadsskjema. Prøv igjen.", "danger")

    # GET (and POST error fallback)
    if "," in user_name:
        parts = user_name.split(",", 1)
        user_name = f"{parts[1].strip()} {parts[0].strip()}"
    default_rullenr_navn = f"Rullenr.: {user_rullenummer} - {user_name}".strip()
    return render_template(
        "søknadsskjema.html",
        page_name="Søknadsskjema",
        favorites=fav_order_lst,
        choices=choices,
        current_turnus_set=user_turnus_set,
        all_turnus_sets=db_utils.get_all_turnus_sets(),
        today=date.today().strftime("%d.%m.%Y"),
        today_iso=date.today().strftime("%Y-%m-%d"),
        default_rullenr_navn=default_rullenr_navn,
        default_stasjoneringssted=user_stasjoneringssted,
    )
